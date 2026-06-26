# app.py 
import streamlit as st
import requests
import json
import pandas as pd
from datetime import datetime
from zoneinfo import ZoneInfo
import io
import re
from collections import Counter
pip install python-docx
# 新增：docx 和 图表 相关库
try:
    import docx
except ImportError:
    docx = None
try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # 避免在Streamlit中显示问题
    from wordcloud import WordCloud
    import numpy as np
except ImportError:
    plt = None
    WordCloud = None
    np = None

# 页面配置
st.set_page_config(
    page_title="GeoField AI - 田野访谈编码助手",
    page_icon="🌍",
    layout="wide"
)

# 标题
st.title("🌍 GeoField AI · 田野访谈编码助手")
st.caption("证据导向型田野资料分析工具 | 支持 TXT / DOCX")

# 初始化session state
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'research_question' not in st.session_state:
    st.session_state.research_question = ""
if 'interview_text' not in st.session_state:
    st.session_state.interview_text = ""
if 'filename' not in st.session_state:
    st.session_state.filename = ""
if 'analysis_data' not in st.session_state:
    st.session_state.analysis_data = None
if 'df' not in st.session_state:
    st.session_state.df = None
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""
if 'has_images' not in st.session_state:
    st.session_state.has_images = False

# 侧边栏 - API配置
with st.sidebar:
    st.header("⚙️ 配置")
    
    # API Key输入
    api_key = st.text_input(
        "DeepSeek API Key",
        type="password",
        help="请输入您的DeepSeek API Key",
        value=st.session_state.api_key
    )
    if api_key:
        st.session_state.api_key = api_key
    
    st.divider()
    
    # 显示当前状态
    st.subheader("📊 进度")
    steps = ["1. 研究问题", "2. 上传访谈", "3. AI分析", "4. 编码审核", "5. 导出报告"]
    current_step = st.session_state.step
    for i, step in enumerate(steps, 1):
        if i < current_step:
            st.success(f"✅ {step}")
        elif i == current_step:
            st.info(f"🔄 {step}")
        else:
            st.text(f"⏳ {step}")

# ==================== 文本提取函数 ====================
def extract_text_from_txt(file):
    """从TXT文件提取文本"""
    try:
        content = file.read().decode('utf-8')
        return content, False
    except UnicodeDecodeError:
        try:
            content = file.read().decode('gbk')
            return content, False
        except:
            raise ValueError("无法解码文件，请确保文件是UTF-8或GBK编码")

def extract_text_from_docx(file):
    """从DOCX文件提取文本（忽略图片）"""
    if docx is None:
        raise ImportError("请安装python-docx库: pip install python-docx")
    
    try:
        doc = docx.Document(io.BytesIO(file.read()))
        
        # 检查是否有图片
        has_images = False
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                has_images = True
                break
        
        # 提取所有段落文本
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        text = '\n'.join(full_text)
        
        return text, has_images
        
    except Exception as e:
        raise ValueError(f"读取DOCX文件失败: {e}")

def extract_text(file, filename):
    """根据文件类型提取文本"""
    if filename.endswith('.txt'):
        return extract_text_from_txt(file)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file)
    else:
        raise ValueError(f"不支持的文件格式: {filename}")

# ==================== 可视化函数 ====================
def create_code_chart(codebook_df):
    """创建编码分类统计图表"""
    if plt is None:
        return None
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    
    # 柱状图
    colors = plt.cm.Set3(range(len(codebook_df)))
    bars = ax1.bar(codebook_df['编码'], codebook_df['频次'], color=colors)
    ax1.set_title('编码频次分布', fontsize=14, fontweight='bold')
    ax1.set_xlabel('编码类别', fontsize=12)
    ax1.set_ylabel('频次', fontsize=12)
    ax1.tick_params(axis='x', rotation=45)
    
    # 在柱子上显示数值
    for bar, value in zip(bars, codebook_df['频次']):
        height = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2., height,
                f'{value}', ha='center', va='bottom', fontsize=10)
    
    # 饼图
    if len(codebook_df) <= 10:  # 饼图适合类别不太多的情况
        wedges, texts, autotexts = ax2.pie(
            codebook_df['频次'], 
            labels=codebook_df['编码'],
            autopct='%1.1f%%',
            colors=colors
        )
        ax2.set_title('编码占比分布', fontsize=14, fontweight='bold')
    else:
        ax2.text(0.5, 0.5, f'编码类别较多\n({len(codebook_df)}类)',
                ha='center', va='center', fontsize=12)
        ax2.set_title('编码占比分布', fontsize=14, fontweight='bold')
    
    plt.tight_layout()
    return fig

def create_wordcloud(text):
    """创建词云图"""
    if WordCloud is None:
        return None
    
    if not text or len(text) < 10:
        return None
    
    try:
        # 简单分词（中文）
        words = re.findall(r'[\u4e00-\u9fff]+', text)
        word_freq = Counter(words)
        
        # 过滤掉单字和常见无意义词
        stopwords = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '来'}
        
        word_freq_filtered = {w: f for w, f in word_freq.items() if len(w) > 1 and w not in stopwords}
        
        if not word_freq_filtered:
            return None
        
        # 生成词云
        wordcloud = WordCloud(
            font_path=None,  # 使用默认字体
            width=800,
            height=400,
            background_color='white',
            max_words=100,
            relative_scaling=0.5,
            colormap='viridis'
        ).generate_from_frequencies(word_freq_filtered)
        
        fig, ax = plt.subplots(figsize=(10, 5))
        ax.imshow(wordcloud, interpolation='bilinear')
        ax.axis('off')
        ax.set_title('访谈文本词云', fontsize=14, fontweight='bold')
        plt.tight_layout()
        
        return fig
        
    except Exception as e:
        st.warning(f"词云生成失败: {e}")
        return None

# ==================== UI 主流程 ====================

# Step 1: 研究问题
st.header("📌 Step 1: 输入研究问题")

research_question = st.text_area(
    "请输入研究问题",
    value=st.session_state.research_question,
    placeholder="例如：本地居民如何感知乡村旅游发展带来的社区变化？",
    height=100
)

if st.button("确认研究问题", type="primary"):
    if research_question.strip():
        st.session_state.research_question = research_question.strip()
        st.session_state.step = 2
        st.rerun()
    else:
        st.warning("请输入研究问题")

# Step 2: 上传访谈文件
if st.session_state.step >= 2:
    st.divider()
    st.header("📂 Step 2: 上传访谈文件")
    
    uploaded_file = st.file_uploader(
        "上传访谈文本文件 (.txt 或 .docx)",
        type=['txt', 'docx'],
        help="支持UTF-8编码的txt文件或docx文件。注意：docx中的图片将被忽略。"
    )
    
    if uploaded_file:
        try:
            filename = uploaded_file.name
            st.session_state.filename = filename
            
            # 提取文本
            text, has_images = extract_text(uploaded_file, filename)
            st.session_state.interview_text = text
            st.session_state.has_images = has_images
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.success(f"✔ 文件导入成功：{filename}")
            with col2:
                st.info(f"📏 字数：{len(text)}")
            with col3:
                if has_images:
                    st.warning("🖼️ 检测到图片，已忽略（AI仅分析文本）")
            
            with st.expander("📄 数据预览（前300字）"):
                st.text(text[:300] + "..." if len(text) > 300 else text)
            
            # 词云预览（可选）
            if st.checkbox("显示访谈文本词云"):
                wordcloud_fig = create_wordcloud(text)
                if wordcloud_fig:
                    st.pyplot(wordcloud_fig)
                else:
                    st.info("文本内容较少或词云生成失败")
            
            if st.button("开始AI分析", type="primary"):
                if not st.session_state.api_key:
                    st.error("⚠️ 请先在侧边栏输入DeepSeek API Key")
                else:
                    st.session_state.step = 3
                    st.rerun()
                    
        except Exception as e:
            st.error(f"文件读取失败：{e}")

# Step 3: AI分析（保持不变）
if st.session_state.step >= 3:
    st.divider()
    st.header("🧠 Step 3: AI分析")
    
    if not st.session_state.analysis_complete:
        with st.spinner("AI分析中，请稍候..."):
            try:
                # 构建Prompt（同之前）
                prompt = f"""
# GeoField AI V2：证据导向型田野资料分析 Prompt

你是一名具有人文地理学、民族志研究和定性研究经验的研究助手。

请根据研究问题分析访谈资料。

分析时必须严格遵守以下原则。

---

【分析原则】

1. 所有分析必须基于访谈原文。

2. 不得推测受访者未明确表达的内容。

3. 不得根据常识补充受访者的动机、情感、价值观和社会背景。

4. 不得为了丰富分析而强行引入理论概念。

5.除非访谈中有明显证据，尽量避免使用抽象理论词汇进行解释，例如：

* 地方感
* 身份认同
* 乡愁
* 空间生产
* 空间抗争
* 底层抵抗
* 社会资本
* 情感依附

以及其他类似的理论性解释。

6. 如果资料不足，请明确写出：

【资料不足，无法判断】

7. 优先保证分析的真实性，而非丰富性。

8. 研究者拥有最终解释权。你的任务是整理材料、发现线索和提示关注点，而不是替代研究者进行理论解释。

---

请严格按照JSON格式输出。

不要输出Markdown。

不要输出解释。

不要输出```json。

不要输出其它文字。

输出必须是一个JSON数组。

数组中的每一个对象代表一条事实编码。

每个对象必须包含以下字段：

id

quote

ai_code

memo_hint

memo_reason

其中：

id：
从1开始编号。

quote：
对应原文。

ai_code：
建议编码。

memo_hint：
若无需Memo，请输出空字符串。

memo_reason：
说明为什么值得记录Memo；若无则输出空字符串。

输出示例：

[
  {{
    "id":1,
    "quote":"……",
    "ai_code":"……",
    "memo_hint":"",
    "memo_reason":""
  }}
]
"""

                prompt += f"""
研究问题：

{st.session_state.research_question}

访谈内容：

{st.session_state.interview_text}

"""

                # 调用API
                headers = {
                    "Authorization": f"Bearer {st.session_state.api_key}",
                    "Content-Type": "application/json"
                }
                
                payload = {
                    "model": "deepseek-chat",
                    "messages": [
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.2
                }
                
                response = requests.post(
                    "https://api.deepseek.com/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=120
                )
                
                if response.status_code == 200:
                    result = response.json()
                    analysis = result["choices"][0]["message"]["content"]
                    
                    # 解析JSON
                    data = json.loads(analysis)
                    st.session_state.analysis_data = data
                    
                    # 创建DataFrame
                    df = pd.DataFrame(data)
                    df = df.rename(columns={
                        "id": "ID",
                        "quote": "原文",
                        "ai_code": "AI建议编码",
                        "memo_hint": "AI Memo提示",
                        "memo_reason": "AI Memo说明",
                    })
                    
                    # 添加研究者填写列
                    df.insert(df.columns.get_loc("AI建议编码") + 1, "研究者编码", "")
                    df["研究者Memo"] = ""
                    
                    st.session_state.df = df
                    st.session_state.analysis_complete = True
                    st.session_state.step = 4
                    st.rerun()
                    
                else:
                    st.error(f"API调用失败 (HTTP {response.status_code})：{response.text}")
                    
            except json.JSONDecodeError as e:
                st.error(f"JSON解析失败：{e}")
                st.text("原始响应内容：")
                st.code(analysis if 'analysis' in locals() else "无响应")
            except Exception as e:
                st.error(f"分析失败：{e}")
    
    else:
        st.success("✅ AI分析已完成！")
        if st.button("进入编码审核"):
            st.session_state.step = 4
            st.rerun()

# Step 4: 编码审核
if st.session_state.step >= 4 and st.session_state.df is not None:
    st.divider()
    st.header("✏️ Step 4: 编码审核")
    
    st.info("💡 提示：在下方表格中修改「研究者编码」和「研究者Memo」，修改会自动保存")
    
    # 显示统计信息
    df = st.session_state.df.copy()
    total_codes = len(df)
    coded_count = df[df["研究者编码"] != ""].shape[0]
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("总编码条目", total_codes)
    with col2:
        st.metric("已审核", coded_count)
    with col3:
        st.metric("待审核", total_codes - coded_count)
    
    st.divider()
    
    # 交互式表格
    st.subheader("📋 编码列表")
    st.caption("点击单元格直接编辑「研究者编码」和「研究者Memo」列")
    
    edited_df = st.data_editor(
        df,
        column_config={
            "ID": st.column_config.NumberColumn("ID", width="small"),
            "原文": st.column_config.TextColumn("原文", width="large"),
            "AI建议编码": st.column_config.TextColumn("AI建议编码", width="medium"),
            "研究者编码": st.column_config.TextColumn(
                "研究者编码 ✏️",
                width="medium",
                help="请输入您修正后的编码"
            ),
            "研究者Memo": st.column_config.TextColumn(
                "研究者Memo ✏️",
                width="large",
                help="记录您的分析思考"
            ),
            "AI Memo提示": st.column_config.TextColumn("AI Memo提示", width="medium"),
            "AI Memo说明": st.column_config.TextColumn("AI Memo说明", width="medium"),
        },
        hide_index=True,
        use_container_width=True,
        num_rows="dynamic"
    )
    
    st.session_state.df = edited_df
    
    # 可视化区域
    st.divider()
    st.subheader("📊 可视化分析")
    
    # 检查是否有编码
    codebook_df = edited_df[edited_df["研究者编码"] != ""].copy()
    
    if len(codebook_df) > 0:
        # 统计编码
        codebook = (
            codebook_df["研究者编码"]
            .value_counts()
            .reset_index()
        )
        codebook.columns = ["编码", "频次"]
        codebook["占比"] = (codebook["频次"] / codebook["频次"].sum() * 100).round(1).astype(str) + "%"
        
        # 显示Codebook
        st.subheader("📊 Codebook预览")
        st.dataframe(codebook, use_container_width=True, hide_index=True)
        
        # 图表
        col1, col2 = st.columns(2)
        with col1:
            fig = create_code_chart(codebook)
            if fig:
                st.pyplot(fig)
            else:
                st.info("图表生成失败")
        
        with col2:
            # 词云（基于已编码的原文）
            coded_text = " ".join(codebook_df["原文"].tolist())
            wordcloud_fig = create_wordcloud(coded_text)
            if wordcloud_fig:
                st.pyplot(wordcloud_fig)
            else:
                st.info("词云生成失败")
    else:
        st.info("暂无编码，请在表格中填写「研究者编码」后查看可视化")
    
    # 导出按钮
    st.divider()
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📥 导出CSV编码表", type="secondary"):
            csv = edited_df.to_csv(index=False, encoding="utf-8-sig")
            st.download_button(
                label="点击下载CSV",
                data=csv,
                file_name=f"GeoFieldAI_coding_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv"
            )
    
    with col2:
        if st.button("📄 生成正式报告", type="primary"):
            st.session_state.step = 5
            st.rerun()

# Step 5: 生成报告
if st.session_state.step >= 5 and st.session_state.df is not None:
    st.divider()
    st.header("📄 Step 5: 正式报告")
    
    df_final = st.session_state.df.copy()
    
    with st.spinner("正在生成报告..."):
        codebook_df = df_final[df_final["研究者编码"] != ""].copy()
        
        if len(codebook_df) > 0:
            codebook = (
                codebook_df["研究者编码"]
                .value_counts()
                .reset_index()
            )
            codebook.columns = ["编码", "频次"]
            codebook["占比"] = (codebook["频次"] / codebook["频次"].sum() * 100).round(1).astype(str) + "%"
            
            # 显示统计图表
            st.subheader("📊 分析统计")
            col1, col2 = st.columns(2)
            with col1:
                fig = create_code_chart(codebook)
                if fig:
                    st.pyplot(fig)
            with col2:
                # 访谈词云
                wordcloud_fig = create_wordcloud(st.session_state.interview_text)
                if wordcloud_fig:
                    st.pyplot(wordcloud_fig)
            
            st.divider()
            
            # 显示Codebook
            st.subheader("📊 Codebook")
            st.dataframe(codebook, use_container_width=True, hide_index=True)
            
            # 生成Markdown报告
            codebook_md = codebook.to_markdown(index=False)
            
            rows_md = ""
            for _, row in df_final.iterrows():
                rows_md += f"**[{row['ID']}]** {row['原文']}\n\n"
                rows_md += f"- AI建议编码：{row['AI建议编码']}\n"
                rows_md += f"- 研究者编码：{row['研究者编码'] or '（未填写）'}\n"
                rows_md += f"- Memo：{row['研究者Memo'] if str(row['研究者Memo']) not in ['', 'nan'] else '（无）'}\n\n---\n\n"
            
            report_md = f"""# GeoField AI 分析报告

**生成时间：** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**研究问题：** {st.session_state.research_question}
**编码条目数：** {len(df_final)}

---

## Codebook

{codebook_md}

---

## 逐条编码记录

{rows_md}

---

*本报告由 GeoField AI 生成，最终解释权属于研究者。*
"""
            
            st.subheader("📄 报告预览")
            st.markdown(report_md)
            
            # 下载按钮
            st.download_button(
                label="📥 下载完整报告 (Markdown)",
                data=report_md,
                file_name=f"GeoFieldAI_report_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                mime="text/markdown",
                type="primary"
            )
            
            st.divider()
            if st.button("🔄 重新开始"):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()
                
        else:
            st.warning("⚠️ 请先填写研究者编码后再生成报告")
            if st.button("返回编码审核"):
                st.session_state.step = 4
                st.rerun()

# 页脚
st.divider()
st.caption("🌍 GeoField AI v0.2 · 支持 TXT/DOCX · 词云与可视化")

# ==================== 额外功能：批量导出 ====================
# （可选，如果你需要）
